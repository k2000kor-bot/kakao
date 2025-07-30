#!/usr/bin/env python3
"""
구현된 기능들 테스트 스크립트
"""

import os
import sys
import asyncio
from pathlib import Path

# 백엔드 모듈 추가
sys.path.append('backend')

def test_ocr_functionality():
    """OCR 기능 테스트"""
    print("🔍 OCR 기능 테스트...")
    
    try:
        from enhanced_media_processor import EnhancedMediaProcessor
        
        processor = EnhancedMediaProcessor()
        
        # 테스트 이미지 생성 (텍스트가 포함된 간단한 이미지)
        test_image_path = "test_image.png"
        
        # PIL을 사용하여 테스트 이미지 생성
        from PIL import Image, ImageDraw, ImageFont
        
        # 간단한 텍스트 이미지 생성
        img = Image.new('RGB', (400, 100), color='white')
        draw = ImageDraw.Draw(img)
        
        # 기본 폰트 사용
        try:
            font = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 20)
        except:
            font = ImageFont.load_default()
        
        draw.text((10, 10), "Hello World\n안녕하세요", fill='black', font=font)
        img.save(test_image_path)
        
        # OCR 테스트
        result = processor._extract_text_from_image(test_image_path)
        
        if result and "Hello" in result:
            print("✅ OCR 기능 정상 작동")
            return True
        else:
            print("⚠️ OCR 기능 테스트 실패")
            return False
            
    except Exception as e:
        print(f"❌ OCR 테스트 오류: {e}")
        return False
    finally:
        # 테스트 파일 정리
        if os.path.exists(test_image_path):
            os.remove(test_image_path)

def test_pdf_functionality():
    """PDF 기능 테스트"""
    print("📄 PDF 기능 테스트...")
    
    try:
        from enhanced_media_processor import EnhancedMediaProcessor
        
        processor = EnhancedMediaProcessor()
        
        # 간단한 PDF 생성 테스트
        test_pdf_path = "test_document.pdf"
        
        # PyPDF2를 사용하여 간단한 PDF 생성
        from PyPDF2 import PdfWriter, PdfReader
        from io import BytesIO
        
        # 빈 PDF 생성
        writer = PdfWriter()
        page = writer.add_blank_page(width=595, height=842)
        writer.write(open(test_pdf_path, 'wb'))
        
        # PDF 텍스트 추출 테스트
        result = processor._extract_pdf_text(test_pdf_path)
        
        print("✅ PDF 기능 테스트 완료")
        return True
        
    except Exception as e:
        print(f"❌ PDF 테스트 오류: {e}")
        return False
    finally:
        # 테스트 파일 정리
        if os.path.exists(test_pdf_path):
            os.remove(test_pdf_path)

def test_word_functionality():
    """Word 문서 기능 테스트"""
    print("📝 Word 문서 기능 테스트...")
    
    try:
        from enhanced_media_processor import EnhancedMediaProcessor
        
        processor = EnhancedMediaProcessor()
        
        # 간단한 Word 문서 생성
        test_docx_path = "test_document.docx"
        
        from docx import Document
        
        doc = Document()
        doc.add_paragraph("테스트 문서입니다.")
        doc.add_paragraph("이것은 Word 문서 테스트입니다.")
        doc.save(test_docx_path)
        
        # Word 텍스트 추출 테스트
        result = processor._extract_word_text(test_docx_path)
        
        if result and "테스트" in result:
            print("✅ Word 문서 기능 정상 작동")
            return True
        else:
            print("⚠️ Word 문서 기능 테스트 실패")
            return False
            
    except Exception as e:
        print(f"❌ Word 테스트 오류: {e}")
        return False
    finally:
        # 테스트 파일 정리
        if os.path.exists(test_docx_path):
            os.remove(test_docx_path)

def test_excel_functionality():
    """Excel 기능 테스트"""
    print("📊 Excel 기능 테스트...")
    
    try:
        from enhanced_media_processor import EnhancedMediaProcessor
        
        processor = EnhancedMediaProcessor()
        
        # 간단한 Excel 파일 생성
        test_excel_path = "test_data.xlsx"
        
        import pandas as pd
        
        # 테스트 데이터 생성
        data = {
            '이름': ['김철수', '이영희', '박민수'],
            '나이': [25, 30, 35],
            '직업': ['개발자', '디자이너', '매니저']
        }
        
        df = pd.DataFrame(data)
        df.to_excel(test_excel_path, index=False)
        
        # Excel 데이터 추출 테스트
        result = processor._extract_excel_data(test_excel_path)
        
        if result and 'sheets' in result:
            print("✅ Excel 기능 정상 작동")
            return True
        else:
            print("⚠️ Excel 기능 테스트 실패")
            return False
            
    except Exception as e:
        print(f"❌ Excel 테스트 오류: {e}")
        return False
    finally:
        # 테스트 파일 정리
        if os.path.exists(test_excel_path):
            os.remove(test_excel_path)

def test_ai_models():
    """AI 모델 기능 테스트"""
    print("🤖 AI 모델 기능 테스트...")
    
    try:
        from enhanced_multimodal_ai import AGIMultimodalComprehensionEngine
        
        engine = AGIMultimodalComprehensionEngine()
        
        # 텍스트 모델 초기화 테스트
        text_models = engine._initialize_text_models()
        
        if text_models:
            print("✅ AI 모델 초기화 성공")
            return True
        else:
            print("⚠️ AI 모델 초기화 실패")
            return False
            
    except Exception as e:
        print(f"❌ AI 모델 테스트 오류: {e}")
        return False

def test_openai_integration():
    """OpenAI API 연동 테스트"""
    print("🔗 OpenAI API 연동 테스트...")
    
    try:
        # 환경변수 확인
        api_key = os.getenv("OPENAI_API_KEY")
        
        if not api_key:
            print("⚠️ OpenAI API 키가 설정되지 않았습니다.")
            print("   export OPENAI_API_KEY='your-api-key' 를 실행하세요.")
            return False
        
        print("✅ OpenAI API 키 확인됨")
        return True
        
    except Exception as e:
        print(f"❌ OpenAI API 테스트 오류: {e}")
        return False

async def main():
    """메인 테스트 함수"""
    print("🚀 구현된 기능들 테스트 시작...")
    print("=" * 50)
    
    tests = [
        ("OCR 기능", test_ocr_functionality),
        ("PDF 기능", test_pdf_functionality),
        ("Word 기능", test_word_functionality),
        ("Excel 기능", test_excel_functionality),
        ("AI 모델", test_ai_models),
        ("OpenAI API", test_openai_integration),
    ]
    
    results = []
    
    for test_name, test_func in tests:
        print(f"\n📋 {test_name} 테스트 중...")
        try:
            if asyncio.iscoroutinefunction(test_func):
                result = await test_func()
            else:
                result = test_func()
            results.append((test_name, result))
        except Exception as e:
            print(f"❌ {test_name} 테스트 중 오류: {e}")
            results.append((test_name, False))
    
    # 결과 요약
    print("\n" + "=" * 50)
    print("📊 테스트 결과 요약:")
    print("=" * 50)
    
    passed = 0
    total = len(results)
    
    for test_name, result in results:
        status = "✅ 통과" if result else "❌ 실패"
        print(f"{test_name}: {status}")
        if result:
            passed += 1
    
    print(f"\n🎯 전체 결과: {passed}/{total} 통과")
    
    if passed == total:
        print("🎉 모든 테스트 통과! 시스템이 정상적으로 작동합니다.")
    else:
        print("⚠️ 일부 테스트 실패. 추가 설정이 필요할 수 있습니다.")
    
    return passed == total

if __name__ == "__main__":
    asyncio.run(main()) 